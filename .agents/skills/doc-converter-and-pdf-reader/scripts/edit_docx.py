# /// script
# requires-python = ">=3.10"
# dependencies = [
#     "python-docx",
# ]
# ///

import sys
import os
import docx

def replace_in_docx(input_path, output_path, search_str, replace_str):
    if not os.path.exists(input_path):
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        return
        
    doc = docx.Document(input_path)
    
    def process_paragraphs(paragraphs):
        for p in paragraphs:
            if search_str in p.text:
                found_in_run = False
                for run in p.runs:
                    if search_str in run.text:
                        run.text = run.text.replace(search_str, replace_str)
                        found_in_run = True
                
                # Fallback: if search text exists in paragraph but wasn't fully contained in a single run
                if not found_in_run:
                    p.text = p.text.replace(search_str, replace_str)
                     
    # Main body paragraphs
    process_paragraphs(doc.paragraphs)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
                
    # Headers and Footers
    for section in doc.sections:
        if section.header:
            process_paragraphs(section.header.paragraphs)
        if section.footer:
            process_paragraphs(section.footer.paragraphs)
            
    doc.save(output_path)
    print(f"Successfully replaced '{search_str}' with '{replace_str}' in {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 5:
        print("Usage: uv run edit_docx.py <input.docx> <output.docx> <search_text> <replacement_text>")
        sys.exit(1)
    replace_in_docx(sys.argv[1], sys.argv[2], sys.argv[3], sys.argv[4])
        
