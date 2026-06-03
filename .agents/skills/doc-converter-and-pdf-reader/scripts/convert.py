# /// script
# requires-python = ">=3.10"
# dependencies = [
#     "python-docx",
#     "pillow",
# ]
# ///

import os
import sys
import re
from PIL import Image
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def get_png_dimensions(file_path):
    try:
        with Image.open(file_path) as img:
            return img.size
    except Exception as e:
        print(f"Error reading image dimensions: {e}", file=sys.stderr)
        return None

MAX_WIDTH_INCHES = 5.7
def get_scaled_dimensions(file_path):
    dims = get_png_dimensions(file_path)
    if not dims:
        return 5.2, 3.5
    
    width_px, height_px = dims
    width_in = width_px / 96.0
    height_in = height_px / 96.0
    
    if width_in > MAX_WIDTH_INCHES:
        ratio = MAX_WIDTH_INCHES / width_in
        width_in = MAX_WIDTH_INCHES
        height_in = height_in * ratio
        
    return width_in, height_in

def set_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    pPr.append(shd)

def set_paragraph_borders(paragraph, left_color="CCCCCC", size=12):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="{size}" w:space="8" w:color="{left_color}"/>'
        f'<w:top w:val="single" w:sz="4" w:space="8" w:color="E5E5E5"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="8" w:color="E5E5E5"/>'
        f'<w:right w:val="single" w:sz="4" w:space="8" w:color="E5E5E5"/>'
        f'</w:pBdr>'
    )
    pPr.append(pbdr)

def parse_inline(paragraph, text):
    pattern = re.compile(r'(\*\*|__|\*|_|`)')
    parts = pattern.split(text)
    
    bold = False
    italic = False
    code = False
    
    for part in parts:
        if not part:
            continue
        if part in ('**', '__'):
            bold = not bold
            continue
        elif part in ('*', '_'):
            italic = not italic
            continue
        elif part in ('`',):
            code = not code
            continue
            
        run = paragraph.add_run(part)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xA3, 0x15, 0x15)
            rPr = run._r.get_or_add_rPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F4F4F4"/>')
            rPr.append(shd)

def add_code_block(doc, code_lines):
    code_text = "\n".join(code_lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    set_paragraph_shading(p, "F4F4F5")
    set_paragraph_borders(p, left_color="CCCCCC", size=12)

def run_conversion():
    if len(sys.argv) < 3:
        print("Usage: uv run convert.py <input_md_file> <output_docx_file>")
        sys.exit(1)
        
    md_file_path = os.path.abspath(sys.argv[1])
    output_docx_path = os.path.abspath(sys.argv[2])
    md_dir = os.path.dirname(md_file_path)
    
    if not os.path.exists(md_file_path):
        print(f"Error: File not found: {md_file_path}", file=sys.stderr)
        sys.exit(1)
        
    with open(md_file_path, "r", encoding="utf-8") as f:
        md_content = f.read()
        
    lines = md_content.splitlines()
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    in_code_block = False
    code_lines = []
    
    for line in lines:
        # Handle code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block - join with newlines and add to single paragraph
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_lines.append(line)
            continue
            
        # Horizontal rules
        if line.strip() == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            pPr = p._p.get_or_add_pPr()
            pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="8" w:color="E5E5E5"/></w:pBdr>')
            pPr.append(pbdr)
            continue
            
        # Empty lines
        if line.strip() == "":
            continue
            
        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            
            if level == 1:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(text)
                run.bold = True
                run.font.name = 'Segoe UI'
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(text)
                run.bold = True
                run.font.name = 'Segoe UI'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0x00, 0x78, 0xD4)
            else:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(5)
                run = p.add_run(text)
                run.bold = True
                run.font.name = 'Segoe UI'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
            continue
            
        # Images
        img_match = re.match(r'^\s*!\[(.*?)\]\((.*?)\)\s*$', line)
        if img_match:
            alt = img_match.group(1)
            img_path = img_match.group(2)
            full_img_path = os.path.abspath(os.path.join(md_dir, img_path))
            
            if os.path.exists(full_img_path):
                w_in, h_in = get_scaled_dimensions(full_img_path)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                
                run = p.add_run()
                run.add_picture(full_img_path, width=Inches(w_in), height=Inches(h_in))
                
                # Caption
                p_cap = doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(2)
                p_cap.paragraph_format.space_after = Pt(10)
                p_cap.paragraph_format.keep_with_next = False
                
                run_cap = p_cap.add_run(f"Ilustración: {alt}")
                run_cap.font.name = 'Segoe UI'
                run_cap.font.size = Pt(9)
                run_cap.italic = True
                run_cap.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Imagen no encontrada: {img_path}]")
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            continue
            
        # Bullet lists
        bullet_match = re.match(r'^(\s*)[\*\+-]\s+(.*)$', line)
        if bullet_match:
            indent_space = len(bullet_match.group(1))
            content = bullet_match.group(2)
            
            style = 'List Bullet 2' if indent_space > 0 else 'List Bullet'
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            parse_inline(p, content)
            continue
            
        # Paragraphs
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        parse_inline(p, line)
        
    # Flush remaining unclosed code blocks to prevent data loss
    if in_code_block and code_lines:
        add_code_block(doc, code_lines)
        
    doc.save(output_docx_path)
    print(f"Document successfully written to {output_docx_path}")

if __name__ == "__main__":
    run_conversion()
